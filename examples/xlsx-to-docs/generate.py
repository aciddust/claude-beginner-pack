import re
from pathlib import Path
import openpyxl
from docx import Document

DATA_PATH = Path("평가데이터.xlsx")
TEMPLATE_PATH = Path("템플릿.docx")
OUT_DIR = Path("결과")

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def load_rows(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        raise SystemExit("엑셀에 데이터가 없습니다.")
    headers = [str(h).strip() if h is not None else "" for h in rows[0]]
    if not any(headers):
        raise SystemExit("엑셀의 헤더 행이 비어 있습니다.")
    data = []
    for row in rows[1:]:
        if all(v is None for v in row):
            continue
        record = {headers[i]: row[i] for i in range(len(headers)) if headers[i]}
        data.append(record)
    return headers, data


def replace_in_paragraph(paragraph, mapping):
    if "{{" not in paragraph.text:
        return
    runs = paragraph.runs
    if not runs:
        return
    full = "".join(r.text for r in runs)

    def sub(match):
        key = match.group(1).strip()
        value = mapping.get(key, match.group(0))
        return "" if value is None else str(value)

    new_text = PLACEHOLDER_RE.sub(sub, full)
    if new_text == full:
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p, mapping)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p, mapping)


def main():
    headers, data = load_rows(DATA_PATH)
    if not data:
        raise SystemExit("데이터 행이 없습니다.")

    OUT_DIR.mkdir(exist_ok=True)
    template_stem = TEMPLATE_PATH.stem
    template_suffix = TEMPLATE_PATH.suffix

    created, failed = [], []
    for idx, record in enumerate(data, start=1):
        name = record.get("이름") or f"행{idx}"
        out_path = OUT_DIR / f"{name}_{template_stem}{template_suffix}"
        try:
            doc = Document(TEMPLATE_PATH)
            replace_everywhere(doc, record)
            doc.save(out_path)
            created.append(out_path.name)
            print(f"  생성: {out_path}")
        except Exception as e:
            failed.append((name, str(e)))
            print(f"  실패: {name} - {e}")

    print(f"\n완료: {len(created)}건 생성, {len(failed)}건 실패")
    if failed:
        for n, err in failed:
            print(f"  - {n}: {err}")


if __name__ == "__main__":
    main()
